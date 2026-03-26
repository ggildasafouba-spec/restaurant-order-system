from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

out = Path('/mnt/data/restaurant-order-system/docs/dossier-projet.docx')
doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.75)
sec.bottom_margin = Inches(0.75)
sec.left_margin = Inches(0.8)
sec.right_margin = Inches(0.8)

styles = doc.styles
styles['Normal'].font.name = 'Aptos'
styles['Normal'].font.size = Pt(10.5)
styles['Title'].font.name = 'Aptos Display'
styles['Title'].font.size = Pt(24)
styles['Title'].font.bold = True
styles['Heading 1'].font.name = 'Aptos Display'
styles['Heading 1'].font.size = Pt(15)
styles['Heading 1'].font.bold = True
styles['Heading 1'].font.color.rgb = RGBColor(138, 21, 56)
styles['Heading 2'].font.name = 'Aptos'
styles['Heading 2'].font.size = Pt(12)
styles['Heading 2'].font.bold = True
styles['Heading 2'].font.color.rgb = RGBColor(55, 65, 81)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('DOSSIER DE PROJET\n')
r.bold = True; r.font.size = Pt(26); r.font.name = 'Aptos Display'; r.font.color.rgb = RGBColor(138,21,56)
r = p.add_run('Application de gestion de commandes pour restaurant')
r.font.size = Pt(16); r.font.name = 'Aptos'

def spacer(n=1):
    for _ in range(n):
        doc.add_paragraph('')
spacer(2)

info = doc.add_table(rows=4, cols=2)
info.style = 'Table Grid'
labels = ['Projet', 'Architecture', 'Interfaces', 'Technologies']
vals = ['Système de commandes restaurant', 'Serveur Node.js central + clients React', 'Client / Caisse, Cuisinier, Serveur, Suivi client', 'React, Node.js, Express, Socket.io, MongoDB']
for i, (a,b) in enumerate(zip(labels, vals)):
    info.cell(i,0).text = a
    info.cell(i,1).text = b

spacer()
para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = para.add_run("Document préparé pour une soutenance de type TP DWWM. Il décrit le besoin métier, l'architecture, les fonctionnalités et les compétences mobilisées.")
run.italic = True

sections = [
    ('1. Contexte et objectif', [
        "Le projet vise à centraliser la gestion des commandes d'un restaurant dans un système unique accessible par plusieurs rôles.",
        "L'objectif principal est de fluidifier la communication entre la caisse, la cuisine, le service et le client final grâce à une mise à jour en temps réel des statuts de commande.",
    ]),
    ('2. Besoin métier', [
        "Dans un restaurant, les erreurs de transmission et le manque de visibilité sur l'état des commandes ralentissent le service.",
        "La solution proposée permet d'enregistrer une commande, de la faire progresser par étapes normalisées et de rendre chaque changement visible immédiatement sur les interfaces concernées.",
    ]),
    ('3. Acteurs et rôles', []),
]
for title, paras in sections:
    doc.add_heading(title, level=1)
    for txt in paras:
        doc.add_paragraph(txt)

roles = doc.add_table(rows=1, cols=3)
roles.style = 'Table Grid'
hdr = roles.rows[0].cells
hdr[0].text = 'Rôle'; hdr[1].text = 'Responsabilité'; hdr[2].text = 'Interface'
for row in [
    ('Client / Caisse', 'Créer la commande et simuler le paiement', 'Application React dédiée'),
    ('Cuisinier', 'Passer une commande en préparation puis prête', 'Application React cuisine'),
    ('Serveur', 'Voir les commandes prêtes et marquer servie', 'Application React service'),
    ('Client (suivi)', 'Consulter l’avancement de la commande', 'Interface de suivi temps réel'),
]:
    cells = roles.add_row().cells
    for i,v in enumerate(row): cells[i].text = v

doc.add_heading('4. Fonctionnalités principales', level=1)
for heading, items in [
    ('Client / Caisse', ['affichage du menu', 'ajout au panier', 'calcul automatique du total', 'validation de la commande', 'visualisation du suivi']),
    ('Cuisinier', ['réception instantanée des nouvelles commandes', 'passage du statut à en préparation', 'passage du statut à prête']),
    ('Serveur', ['vue de toutes les commandes', 'mise en avant des commandes prêtes', 'passage du statut à servie']),
    ('Suivi client', ['recherche par référence', 'frise d’avancement', 'rafraîchissement automatique sans rechargement']),
]:
    doc.add_heading(heading, level=2)
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

doc.add_heading('5. Architecture technique', level=1)
doc.add_paragraph("L'architecture repose sur un backend Node.js + Express qui expose des routes REST pour le menu et les commandes. Socket.io assure la diffusion immédiate des événements métier vers toutes les interfaces connectées. MongoDB est prévu pour la persistance, avec un mode mémoire de secours pour la démonstration locale.")
arch = doc.add_table(rows=1, cols=3)
arch.style = 'Table Grid'
arch.rows[0].cells[0].text = 'Couche'; arch.rows[0].cells[1].text = 'Choix'; arch.rows[0].cells[2].text = 'Rôle'
for row in [
    ('Front-end', 'React + Vite', 'Créer des interfaces distinctes par métier'),
    ('Back-end', 'Node.js + Express', 'Porter la logique métier et les API'),
    ('Temps réel', 'Socket.io', 'Propager les changements de statut'),
    ('Données', 'MongoDB + Mongoose', 'Structurer menu, commandes et utilisateurs'),
]:
    c = arch.add_row().cells
    for i,v in enumerate(row): c[i].text = v

doc.add_heading('6. Cycle de vie des commandes', level=1)
doc.add_paragraph('Le flux métier partagé entre toutes les interfaces suit quatre statuts simples et compréhensibles :')
flow = doc.add_paragraph()
flow.alignment = WD_ALIGN_PARAGRAPH.CENTER
r=flow.add_run('new  →  in_preparation  →  ready  →  served')
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RGBColor(22, 101, 52)
doc.add_paragraph("Chaque mise à jour déclenche un événement Socket.io afin que l’ensemble des écrans reflète immédiatement l’état courant de la commande.")

doc.add_heading('7. Compétences TP DWWM mobilisées', level=1)
comp = doc.add_table(rows=1, cols=2)
comp.style='Table Grid'
comp.rows[0].cells[0].text='Bloc de compétences'
comp.rows[0].cells[1].text='Illustration dans le projet'
for row in [
    ('Front-end', 'Maquettage des interfaces, réalisation d’écrans statiques, comportements dynamiques React'),
    ('Back-end', 'Conception de modèles, routes API, contrôleurs, composants métier côté serveur'),
    ('Qualité de code', 'Séparation des couches, réutilisation de modules partagés, lisibilité et maintenabilité'),
]:
    c=comp.add_row().cells
    for i,v in enumerate(row): c[i].text=v

doc.add_heading('8. Conclusion', level=1)
doc.add_paragraph("Le projet répond au cahier des charges fonctionnel transmis : il met en scène plusieurs rôles, un suivi temps réel des commandes et une architecture cohérente pour une présentation de fin de parcours. Il peut être complété ensuite par une authentification, une gestion d'administration et un déploiement cloud.")

doc.save(out)
print(out)
